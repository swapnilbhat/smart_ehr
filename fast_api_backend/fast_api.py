from fastapi import FastAPI,Request,HTTPException,File, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel
import motor.motor_asyncio
from bson import ObjectId
import uvicorn
import aiohttp
import openai
import json
import os
import re
import uuid
import time
import datetime
from openai import OpenAI
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from fastapi.middleware.cors import CORSMiddleware
from pdf2image import convert_from_bytes
import pytesseract
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT  # Import vertical and table alignments
from docx.shared import RGBColor
from PIL import Image
from pypandoc import convert_file

app = FastAPI()

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust this to your specific needs
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# MongoDB connection
MONGO_DETAILS = "mongodb://localhost:27017"
client = motor.motor_asyncio.AsyncIOMotorClient(MONGO_DETAILS)
database = client.health_records
ehr_collection = database.get_collection("ehr_2")
patient_lookup=database.get_collection("patient_lookup_2")
REPORTS_DIR = 'reports/'
investigation_collection=database.get_collection("investigations")
investigation_lookup=database.get_collection("investigation_lookup")
patient_id_name_collection=database.get_collection("patient_id_name_mapping")
# Pydantic model for EHR
class EHRModel(BaseModel):
    entry: int
    patient_id: str
    report: dict

class Query(BaseModel):
    query:str

class PatientID(BaseModel):
    patient_id: str
    
# Helper function to convert BSON to JSON
def ehr_helper(ehr) -> dict:
    return {
        "id": str(ehr["_id"]),
        "entry": ehr["entry"],
        "patient_id": ehr["patient_id"],
        "report": ehr["report"]
    }

# Insert a document into the MongoDB collection
#@app.post("/add_record/", response_model=EHRModel)
async def add_record(entry,patient_id,report,isInvestigation=False):
    if isInvestigation:
        now = datetime.datetime.now()
        date_format = now.strftime("%d/%m/%Y")
        time_format = now.strftime("%H:%M")
        report_mod = {
            'Record Entry': {
                'Datetime':now,
                'Date': date_format,
                'Time': time_format
            }
        }
        report_mod.update(report)
        # Create the record to be inserted
        record = {
            'entry': entry,
            'patient_id': patient_id,
            'report': report_mod
        }
        new_record = await investigation_collection.insert_one(record)
        print(f'Record added for patient_id: {patient_id}')
        patient_exists = await patient_id_name_collection.find_one({'patient_id': patient_id})
        print('Patient exists in id name lookup ',patient_exists)
        if not patient_exists:
            # If patient_id doesn't exist, add the entry
            name=search_key_in_json(report,'Name')
            print('Id',patient_id,'Name',name)
            await patient_id_name_collection.insert_one({'patient_id': patient_id, 'name': name})
            print(f'Added new patient entry: patient_id: {patient_id}, name: {name}')
        # Update the lookup table
        lookup_result = await investigation_lookup.update_one(
            {'patient_id': patient_id},
            {'$inc': {'record_count': 1}},
            upsert=True
        )
        created_record = await investigation_collection.find_one({"_id": new_record.inserted_id})
        
        return created_record
    else:
        now = datetime.datetime.now()
        date_format = now.strftime("%d/%m/%Y")
        time_format = now.strftime("%H:%M")
        report_mod = {
            'Record Entry': {
                'Datetime':now,
                'Date': date_format,
                'Time': time_format
            }
        }
        report_mod.update(report)
        # Create the record to be inserted
        record = {
            'entry': entry,
            'patient_id': patient_id,
            'report': report_mod
        }
        new_record = await ehr_collection.insert_one(record)
        print(f'Record added for patient_id: {patient_id}')
        patient_exists = await patient_id_name_collection.find_one({'patient_id': patient_id})
        print('Patient exists in id name lookup ',patient_exists)
        if not patient_exists:
            # If patient_id doesn't exist, add the entry
            name=search_key_in_json(report,'Name')
            print('Id',patient_id,'Name',name)
            await patient_id_name_collection.insert_one({'patient_id': patient_id, 'name': name})
            print(f'Added new patient entry: patient_id: {patient_id}, name: {name}')
        # Update the lookup table
        lookup_result = await patient_lookup.update_one(
            {'patient_id': patient_id},
            {'$inc': {'record_count': 1}},
            upsert=True
        )
        created_record = await ehr_collection.find_one({"_id": new_record.inserted_id})
        
        return created_record

async def get_patient_record_count(patient_id,isInvestigation=False):
    if isInvestigation:
        patient_record = await investigation_lookup.find_one({'patient_id': patient_id})
        if patient_record:
            return patient_record['record_count']
        else:
            return 0
    else:
        patient_record = await patient_lookup.find_one({'patient_id': patient_id})
        if patient_record:
            return patient_record['record_count']
        else:
            return 0

# Register the Unicode font (supports both English and Hindi)
def register_fonts():
    pdfmetrics.registerFont(TTFont('NotoSans', r'C:\Users\swapnil\OneDrive\Documents\ehr_on_github\smart_ehr\fast_api_backend\Noto_Sans\static\NotoSans_Condensed-Regular.ttf'))  # Path to the Unicode font
    pdfmetrics.registerFont(TTFont('NotoSans-Bold', r'C:\Users\swapnil\OneDrive\Documents\ehr_on_github\smart_ehr\fast_api_backend\Noto_Sans\static\NotoSans_Condensed-Bold.ttf'))  # Bold font
    
async def print_report(patient_id,isInvestigation=False):
    register_fonts()  # Register the font at the start
    if isInvestigation:
        #Investigation
        # Retrieve all reports for the given patient_id
        records = await investigation_collection.find({'patient_id': patient_id}).sort('entry').to_list(None)
        # Check if records exist
        if not records:
            print(f'No records found for patient_id: {patient_id}')
            return
        
        print('found records')
        # Create a PDF document
        #HEADER_IMAGE='./kem_logo.png'
        HEADER_IMAGE='./aristalogo.png'
        pdf_file = f'{REPORTS_DIR}{patient_id}_investigation.pdf'
        c = canvas.Canvas(pdf_file, pagesize=letter)
        width, height = letter

        # Define margins
        margin_left = inch
        margin_right = inch
        margin_top = inch
        margin_bottom = inch

        first_entry = True  # Flag to check for the first entry

        for record in records:
            entry = record['entry']
            report = record['report']
            print('print report',report)
            #Remove Datetime key and object
            if 'Datetime' in report['Record Entry']:
                del report['Record Entry']['Datetime']
            print('print report mod',report)
            report = json_to_formatted_string(report)
            
            # Create a new page for each entry
            if not first_entry:
                c.showPage()  # Create a new page for each subsequent entry
            first_entry = False

            # Draw the header
            #c.drawImage(HEADER_IMAGE, margin_left, height - margin_top, width=2*inch, preserveAspectRatio=True, mask='auto')
            try:
                c.drawImage(HEADER_IMAGE, margin_left, height - margin_top -3*inch, width=1*inch, preserveAspectRatio=True, mask='auto')
            except Exception as e:
                print(f"Error loading image: {e}")
            c.setFont("NotoSans-Bold", 14)
            #c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.2 * inch, "Your Report Header Text")
            c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.5 * inch, "Arista Surge Medical Center")
            c.setFont("NotoSans", 12)
            c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.2 * inch, "Mumbai")
            # c.drawString(margin_left + 2.1 * inch, height - margin_top, "Phone: (123) 456-7890")

            # Draw the entry number
            c.setFont("NotoSans", 12)
            c.line(margin_left, height - margin_top - 0.3 * inch, width - margin_right, height - margin_top - 0.3 * inch)
            c.drawString(margin_left, height - margin_top - 0.7 * inch, f"Patient Id: {patient_id}")
            c.drawString(margin_left, height - margin_top - 1* inch, f"Entry Number: {entry}")
            # Print the report content
            y_position = height - margin_top - 1.4*inch
            text_lines = report.split('\n')
            c.setFont("NotoSans", 12)
            
            for line in text_lines:
                # Ensure the text wraps if it's too long
                wrapped_lines = []
                while len(line) > 80:  # 80 characters per line for wrapping
                    wrapped_lines.append(line[:80])
                    line = line[80:]
                wrapped_lines.append(line)
                
                for wrapped_line in wrapped_lines:
                    if y_position < margin_bottom:  # Avoid printing too close to the bottom
                        c.showPage()
                        y_position = height - margin_top - 1.6*inch
                        # Draw the header again
                        #c.drawImage(HEADER_IMAGE, margin_left, height - margin_top, width=2*inch, preserveAspectRatio=True, mask='auto')
                        try:
                            c.drawImage(HEADER_IMAGE, margin_left, height - margin_top -3*inch, width=1*inch, preserveAspectRatio=True, mask='auto')
                        except Exception as e:
                            print(f"Error loading image: {e}")
                        c.setFont("NotoSans-Bold", 14)
                        #c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.2 * inch, "Your Report Header Text")
                        c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.5 * inch, "Arista Surge Medical Center")
                        c.setFont("NotoSans", 12)
                        c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.2 * inch, "Mumbai")
                        # c.drawString(margin_left + 2.1 * inch, height - margin_top, "Phone: (123) 456-7890")
                        c.setFont("NotoSans", 12)
                    #c.drawString(margin_left, y_position, wrapped_line)
                    current_margin_left = margin_left
                    if wrapped_line.startswith("#"):
                        # Heading line
                        y_position -= 10
                        wrapped_line = wrapped_line.replace("#", "").strip()
                        c.setFont("NotoSans-Bold", 14)
                        c.drawString(current_margin_left, y_position, wrapped_line)
                    elif '**' in wrapped_line:
                        # Subheading line
                        parts = wrapped_line.split('**')
                        for i,part in enumerate(parts):
                            if i % 2 == 1:  # Odd indices represent the bold text
                                c.setFont("NotoSans-Bold", 12)
                            else:
                                c.setFont("NotoSans", 12)
                            c.drawString(current_margin_left, y_position, part.strip())
                            current_margin_left += c.stringWidth(part.strip())       
                    else:
                        # Normal text
                        c.setFont("NotoSans", 12)
                        c.drawString(current_margin_left, y_position, wrapped_line)
                    y_position -= 20  # Adjust line spacing for better readability

            # Add a line separator between reports
            # y_position -= 5
            # c.line(margin_left, y_position, width - margin_right, y_position)
            c.line(margin_left, margin_bottom, width - margin_right, margin_bottom)
            y_position -= 25  # Additional spacing before the next entry

            # Draw the footer
            c.setFont("NotoSans", 8)
            c.drawString(margin_left, margin_bottom / 2, f"Page {c.getPageNumber()} - Confidential")

        c.save()
        print(f'Investigation for patient_id {patient_id} has been saved to {pdf_file}')
    else:
        #Report
        # Retrieve all reports for the given patient_id
        records = await ehr_collection.find({'patient_id': patient_id}).sort('entry').to_list(None)
        # Check if records exist
        if not records:
            print(f'No records found for patient_id: {patient_id}')
            return
        
        print('found records')
        # Create a PDF document
        # HEADER_IMAGE='./kem_logo.png'
        HEADER_IMAGE='./aristalogo.png'
        pdf_file = f'{REPORTS_DIR}{patient_id}_report.pdf'
        c = canvas.Canvas(pdf_file, pagesize=letter)
        width, height = letter

        # Define margins
        margin_left = inch
        margin_right = inch
        margin_top = inch
        margin_bottom = inch

        first_entry = True  # Flag to check for the first entry

        for record in records:
            entry = record['entry']
            report = record['report']
            print('print report',report)
            #Remove Datetime key and object
            if 'Datetime' in report['Record Entry']:
                del report['Record Entry']['Datetime']
            print('print report mod',report)
            report = json_to_formatted_string(report)
            
            # Create a new page for each entry
            if not first_entry:
                c.showPage()  # Create a new page for each subsequent entry
            first_entry = False

            # Draw the header
            #c.drawImage(HEADER_IMAGE, margin_left, height - margin_top, width=2*inch, preserveAspectRatio=True, mask='auto')
            # margin_top -1.8*inch
            try:
                c.drawImage(HEADER_IMAGE, margin_left, height - margin_top -3*inch, width=1*inch, preserveAspectRatio=True, mask='auto')
            except Exception as e:
                print(f"Error loading image: {e}")
            c.setFont("NotoSans-Bold", 14)
            #c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.2 * inch, "Your Report Header Text")
            c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.5 * inch, "Arista Surge Medical Center")
            c.setFont("NotoSans", 12)
            c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.2 * inch, "Mumbai")
            # c.drawString(margin_left + 2.1 * inch, height - margin_top, "Phone: (123) 456-7890")

            # Draw the entry number
            c.setFont("NotoSans", 12)
            c.line(margin_left, height - margin_top - 0.3 * inch, width - margin_right, height - margin_top - 0.3 * inch)
            c.drawString(margin_left, height - margin_top - 0.7 * inch, f"Patient Id: {patient_id}")
            c.drawString(margin_left, height - margin_top - 1* inch, f"Entry Number: {entry}")
            # Print the report content
            y_position = height - margin_top - 1.4*inch
            text_lines = report.split('\n')
            c.setFont("NotoSans", 12)
            
            for line in text_lines:
                # Ensure the text wraps if it's too long
                wrapped_lines = []
                while len(line) > 80:  # 80 characters per line for wrapping
                    wrapped_lines.append(line[:80])
                    line = line[80:]
                wrapped_lines.append(line)
                for wrapped_line in wrapped_lines:
                    if y_position < margin_bottom:  # Avoid printing too close to the bottom
                        c.showPage()
                        y_position = height - margin_top - 1.6*inch
                        # Draw the header again
                        #c.drawImage(HEADER_IMAGE, margin_left, height - margin_top, width=2*inch, preserveAspectRatio=True, mask='auto')
                        try:
                            c.drawImage(HEADER_IMAGE, margin_left, height - margin_top -3*inch, width=1*inch, preserveAspectRatio=True, mask='auto')
                        except Exception as e:
                            print(f"Error loading image: {e}")
                        c.setFont("NotoSans-Bold", 14)
                        #c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.2 * inch, "Your Report Header Text")
                        c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.5 * inch, "Arista Surge Medical Center")
                        c.setFont("NotoSans", 12)
                        c.drawString(margin_left + 2.1 * inch, height - margin_top + 0.2 * inch, "Mumbai")
                        # c.drawString(margin_left + 2.1 * inch, height - margin_top, "Phone: (123) 456-7890")
                        c.setFont("NotoSans", 12)
                    #Old code without bold
                    #c.drawString(margin_left, y_position, wrapped_line)
                    
                    #New code with bold
                    # Check for bold markers and apply bold formatting
                    current_margin_left = margin_left
                    # if '**' in wrapped_line:
                    #     parts = wrapped_line.split('**')
                    #     for i, part in enumerate(parts):
                    #         if i % 2 == 1:  # Odd indices represent the bold text
                    #             c.setFont("Helvetica-Bold", 12)
                    #         else:
                    #             c.setFont("Helvetica", 12)
                    #         c.drawString(current_margin_left, y_position, part.strip())
                    #         current_margin_left += c.stringWidth(part.strip())  # Adjust position based on the width of the printed text
                    # else:
                    #     c.setFont("Helvetica", 12)
                    #     c.drawString(current_margin_left, y_position, wrapped_line)
                     # Apply different styles based on heading markers
                    if wrapped_line.startswith("#"):
                        # Heading line
                        y_position -= 10
                        wrapped_line = wrapped_line.replace("#", "").strip()
                        c.setFont("NotoSans-Bold", 14)
                        c.drawString(current_margin_left, y_position, wrapped_line)
                    elif '**' in wrapped_line:
                        # Subheading line
                        parts = wrapped_line.split('**')
                        for i,part in enumerate(parts):
                            if i % 2 == 1:  # Odd indices represent the bold text
                                c.setFont("NotoSans-Bold", 12)
                            else:
                                c.setFont("NotoSans", 12)
                            c.drawString(current_margin_left, y_position, part.strip())
                            current_margin_left += c.stringWidth(part.strip())       
                    else:
                        # Normal text
                        c.setFont("NotoSans", 12)
                        c.drawString(current_margin_left, y_position, wrapped_line)
                        
                    y_position -= 20  # Adjust line spacing for better readability

            # Add a line separator between reports
            # y_position -= 5
            # c.line(margin_left, y_position, width - margin_right, y_position)
            c.line(margin_left, margin_bottom, width - margin_right, margin_bottom)
            y_position -= 30  # Additional spacing before the next entry

            # Draw the footer
            c.setFont("NotoSans", 8)
            c.drawString(margin_left, margin_bottom / 2, f"Page {c.getPageNumber()} - Confidential")

        c.save()
        print(f'Report for patient_id {patient_id} has been saved to {pdf_file}')


# @app.get("/print_word_report/")
# async def print_word_report(patient_id: str): #patient_id,isInvestigation=False
async def print_word_report(patient_id,isInvestigation=False):
    if isInvestigation:
        records = await investigation_collection.find({'patient_id': patient_id}).sort('entry').to_list(None)
        
        # Check if records exist
        if not records:
            print(f'No records found for patient_id: {patient_id}')
            raise HTTPException(status_code=404, detail="No records found for this patient ID")
        
        print('found records')
        doc = Document()
        def set_margins(doc, left=1, right=0.5, top=1.0, bottom=1.0):
            """
            Sets the page margins for the document.
            """
            sections = doc.sections
            for section in sections:
                section.left_margin = Inches(left)
                section.right_margin = Inches(right)
                section.top_margin = Inches(top)
                section.bottom_margin = Inches(bottom)
        set_margins(doc, left=1, right=0.5, top=1.0, bottom=1.0)  # Adjust margins here
        
        # Set document styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Noto Sans'
        font.size = Pt(12)

        HEADER_IMAGE = './aristalogo.png'
        
        def add_header(doc):
            try:
                # Resize the image if necessary before adding it
                img = Image.open(HEADER_IMAGE)
                img.thumbnail((Inches(1), Inches(1)))  # Adjust the image size
                img.save(HEADER_IMAGE)
            except Exception as e:
                print(f"Error loading image: {e}")
            
            # Create a table with one row and two columns
            table = doc.add_table(rows=1, cols=2)

            # Add the image to the left cell
            cell_image = table.cell(0, 0)
            cell_image.width = Inches(1.5)
            paragraph_image = cell_image.paragraphs[0]
            run_image = paragraph_image.add_run()
            run_image.add_picture(HEADER_IMAGE, width=Inches(1.5))  # Adjust width as needed

            # Add the clinic name and address to the right cell
            cell_text = table.cell(0, 1)
            cell_text.width = Inches(5)
            cell_text_paragraph = cell_text.paragraphs[0]
            
            # Add clinic name (Heading 1)
            run_name = cell_text_paragraph.add_run("Arista Surge Medical Center")
            run_name.bold = True
            run_name.font.size = Pt(16)
            
            # Add address (Normal text)
            cell_text_paragraph.add_run("\nMumbai").font.size = Pt(12)

            # Set the alignment for text (centered vertically)
            table.autofit = False
            for cell in table.columns[1].cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Add a horizontal line below the header
            #doc.add_paragraph().add_run().add_break()
            para = doc.add_paragraph()
            para.add_run("_" * 100)  # This creates a horizontal line stretching across the page
            

        first_entry = True
        
        # Iterate through each record
        for record in records:
            if not first_entry:
                doc.add_page_break()  # Add a page break between entries
            
            add_header(doc)  # Add header to each page
            
            entry = record['entry']
            report = record['report']
            print('print report', report)
            
            # Remove Datetime key and object
            if 'Datetime' in report['Record Entry']:
                del report['Record Entry']['Datetime']
            
            print('print report mod', report)
            # Add Patient ID and Entry Number
            doc.add_paragraph(f"Patient Id: {patient_id}")
            doc.add_paragraph(f"Entry Number: {entry}")
            
            # Format the report content
            report_text = json_to_formatted_string(report)
            text_lines = report_text.split('\n')
            
            for line in text_lines:
                if line.startswith("#"):
                    # Heading
                    line = line.replace("#", "").strip()
                    doc.add_paragraph(line, style='Heading 2')
                elif '**' in line:
                    # Bold for text enclosed by ** (subheading)
                    parts = line.split('**')
                    para = doc.add_paragraph()
                    for i, part in enumerate(parts):
                        if i % 2 == 1:
                            run = para.add_run(part)
                            run.bold = True
                        else:
                            para.add_run(part)
                else:
                    # Normal text
                    doc.add_paragraph(line)
            
            # Add a line separator between reports
            doc.add_paragraph('-' * 40)
            
            first_entry = False  # Disable the first entry flag after the first iteration
        
        # Save the document
        word_file = f'{REPORTS_DIR}{patient_id}_investigation.docx'
        doc.save(word_file)
        print(f'Investigation for patient_id {patient_id} has been saved to {word_file}')
        
    else:
        records = await ehr_collection.find({'patient_id': patient_id}).sort('entry').to_list(None)
        
        # Check if records exist
        if not records:
            print(f'No records found for patient_id: {patient_id}')
            raise HTTPException(status_code=404, detail="No records found for this patient ID")
        
        print('found records')
        doc = Document()
        def set_margins(doc, left=1, right=0.5, top=1.0, bottom=1.0):
            """
            Sets the page margins for the document.
            """
            sections = doc.sections
            for section in sections:
                section.left_margin = Inches(left)
                section.right_margin = Inches(right)
                section.top_margin = Inches(top)
                section.bottom_margin = Inches(bottom)
        set_margins(doc, left=1, right=0.5, top=1.0, bottom=1.0)  # Adjust margins here
        
        # Set document styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Noto Sans'
        font.size = Pt(12)

        HEADER_IMAGE = './aristalogo.png'
        
        def add_header(doc):
            try:
                # Resize the image if necessary before adding it
                img = Image.open(HEADER_IMAGE)
                img.thumbnail((Inches(1), Inches(1)))  # Adjust the image size
                img.save(HEADER_IMAGE)
            except Exception as e:
                print(f"Error loading image: {e}")
            
            # Create a table with one row and two columns
            table = doc.add_table(rows=1, cols=2)

            # Add the image to the left cell
            cell_image = table.cell(0, 0)
            cell_image.width = Inches(1.5)
            paragraph_image = cell_image.paragraphs[0]
            run_image = paragraph_image.add_run()
            run_image.add_picture(HEADER_IMAGE, width=Inches(1.5))  # Adjust width as needed

            # Add the clinic name and address to the right cell
            cell_text = table.cell(0, 1)
            cell_text.width = Inches(5)
            cell_text_paragraph = cell_text.paragraphs[0]
            
            # Add clinic name (Heading 1)
            run_name = cell_text_paragraph.add_run("Arista Surge Medical Center")
            run_name.bold = True
            run_name.font.size = Pt(16)
            
            # Add address (Normal text)
            cell_text_paragraph.add_run("\nMumbai").font.size = Pt(12)

            # Set the alignment for text (centered vertically)
            table.autofit = False
            for cell in table.columns[1].cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Add a horizontal line below the header
            #doc.add_paragraph().add_run().add_break()
            para = doc.add_paragraph()
            para.add_run("_" * 100)  # This creates a horizontal line stretching across the page
            

        first_entry = True
        
        # Iterate through each record
        for record in records:
            if not first_entry:
                doc.add_page_break()  # Add a page break between entries
            
            add_header(doc)  # Add header to each page
            
            entry = record['entry']
            report = record['report']
            print('print report', report)
            
            # Remove Datetime key and object
            if 'Datetime' in report['Record Entry']:
                del report['Record Entry']['Datetime']
            
            print('print report mod', report)
            # Add Patient ID and Entry Number
            doc.add_paragraph(f"Patient Id: {patient_id}")
            doc.add_paragraph(f"Entry Number: {entry}")
            
            # Format the report content
            report_text = json_to_formatted_string(report)
            text_lines = report_text.split('\n')
            
            for line in text_lines:
                if line.startswith("#"):
                    # Heading
                    line = line.replace("#", "").strip()
                    doc.add_paragraph(line, style='Heading 2')
                elif '**' in line:
                    # Bold for text enclosed by ** (subheading)
                    parts = line.split('**')
                    para = doc.add_paragraph()
                    for i, part in enumerate(parts):
                        if i % 2 == 1:
                            run = para.add_run(part)
                            run.bold = True
                        else:
                            para.add_run(part)
                else:
                    # Normal text
                    doc.add_paragraph(line)
            
            # Add a line separator between reports
            doc.add_paragraph('-' * 40)
            
            first_entry = False  # Disable the first entry flag after the first iteration
        
        # Save the document
        word_file = f'{REPORTS_DIR}{patient_id}_report.docx'
        doc.save(word_file)
        print(f'Report for patient_id {patient_id} has been saved to {word_file}')
        
        #return {'message': f"Report saved successfully for patient_id {patient_id}", 'file_path': word_file}
    
def search_key_in_json(json_obj, search_term):
    search_term = search_term.lower()
    
    def recursive_search(d):
        for key, value in d.items():
            if key.lower() == search_term:
                return value
            elif isinstance(value, dict):
                result = recursive_search(value)
                if result is not None:
                    return result
        return None
    
    return recursive_search(json_obj)

def delete_key_from_json(data, key_to_delete):
    key_to_delete_lower = key_to_delete.lower()

    # Helper function to search and delete the key
    def recursive_delete(d):
        if isinstance(d, dict):
            keys_to_delete = [key for key in d.keys() if key.lower() == key_to_delete_lower]
            for key in keys_to_delete:
                del d[key]
            for key, value in d.items():
                recursive_delete(value)
        elif isinstance(d, list):
            for item in d:
                recursive_delete(item)
    
    # Make a copy of the data to avoid modifying the original object
    data_copy = data.copy()
    recursive_delete(data_copy)
    
    return data_copy


def json_to_formatted_string(json_obj):
    def parse_dict(d, level=0):
        lines = []
        for key, value in d.items():
            if isinstance(value, dict):
                # Mark heading for top-level keys, and subheading for nested keys
                if level == 0:
                    lines.append(f"{' ' * (level * 2)}#{key}# :")  # Mark heading key with #
                else:
                    lines.append(f"{' ' * (level * 2)}**{key}**:")  # Mark subheading key with *
                
                lines.extend(parse_dict(value, level + 1))
            elif isinstance(value, list):
                if level == 0:
                    lines.append(f"{' ' * (level * 2)}#{key}#:")  # Mark heading key with #
                else:
                    lines.append(f"{' ' * (level * 2)}**{key}**:")  # Mark subheading key with *
                
                for item in value:
                    if isinstance(item, dict):
                        lines.extend(parse_dict(item, level + 1))
                    else:
                        lines.append(f"{' ' * ((level + 1) * 2)}- {item}")
            else:
                # Mark leaf nodes (key-value pairs)
                if level == 0:
                    lines.append(f"{' ' * (level * 2)}- #{key}#: {value}")  # Heading for level 0
                else:
                    lines.append(f"{' ' * (level * 2)}- **{key}**: {value}")  # Subheading for nested keys
        return lines
    
    formatted_string = "\n".join(parse_dict(json_obj))
    return formatted_string

            
async def gpt_json(prompt,max_tokens):
    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
    messages = [
           {"role": "system", "content": "You are a helpful assistant designed to output JSON."},
            {"role": "user", "content":  f'''{prompt}'''},
            ]
    response = client.chat.completions.create(
        model="gpt-4o-mini", #gpt-4o-mini, gpt-3.5-turbo-0125
        temperature=0,
        response_format={ "type": "json_object" },
        messages=messages
        )
    result=response.choices[0].message.content
    return result

async def task_on_EHR(reports,datetime_now,task):
    task_prompt=f"""
    You are an AI designed to assist doctors in analyzing Electronic Health Records (EHRs). You will be provided with a list of EHRs in json format for a specific patient and a task to perform on these records. The records are presented in chronological order based on their entry number. A lower entry number denotes an earlier record. You are also provided the current datatime, which you can use to compare with the Date and Time mentioned in the EHR.
    The output of the task must be to the point and concise. It must be in the form of a paragraph with numbering(or bullets) and points and spaces wherever required.

    List of EHRs:
    {reports}
    
    Current Datetime:
    {datetime_now}
    
    Task to be perfomed on the EHRs:
    {task}
    
    Your output must be in json format as follows:
    output:<output of task on EHRs>
    """
    output=await gpt_json(task_prompt,1000)
    output=json.loads(output)
    #print(output)
    return output


async def get_latest_patient_id():
    latest_patient = await patient_id_name_collection.find_one(sort=[("_id", -1)])  # Sort by _id descending
    if latest_patient:
        current_id = latest_patient.get("patient_id")
        # Increment the numeric part of the ID
        new_id = int(current_id) + 1
        # Convert back to zero-padded string, keeping the same length as the original
        new_id_str = str(new_id).zfill(len(current_id))  # Zero pad to the same length
        return new_id_str
    else:
        new_id_str='000001'
        return new_id_str
    
    
async def extract_intent_and_content(query:str,intent:str):
    if intent.lower() == 'create':
        create_prompt=f'''You are an AI radiology assistant designed to convert a radiologist’s unstructured report into a well-structured, professional radiology report. The report must adhere strictly to ACR guidelines and be formatted clearly, with appropriate headings and subheadings. Your task is to identify the relevant information, and organize it into a structured format as given under "Radiology Report Structure".
You must strictly include only those Headings/Subheadings whose information is given.
        
##Radiology Report Structure##

Heading: Demographics

Subheading:
Facility/Location: Where the study was performed.

Heading: Patient Information

Subheadings:
Patient Id: Patient Id assigned to the patient
Name: Name of the patient
Age or Date of Birth: Write Age of the patient or Date of Birth of patient, whichever is provided.
Gender: Gender of the patient

Heading: Examination details

Subheadings:
Referring Physician: Name(s) of the referring physician(s) or other healthcare provider(s). Indicate if the patient is self-referred.
Examination Type: Name or type of examination performed.
Examination Date: The date when the examination took place.
Examination Time: Include the time of the examination if relevant.
Additional Information: Optional fields like Date of Dictation and Date/Time of Transcription.

Heading: Relevant Clinical Information

Subheading:
Clinical History: Provide a summary of the patient’s relevant clinical history, including chief complaints and any pertinent medical background that led to the imaging study.

Heading: Procedures and Materials

Subheadings:
Study/Procedure Details: A description of the studies or procedures performed during the examination.
Contrast Media: Details on any contrast media or radiopharmaceuticals used, including dosage, concentration, volume, and route of administration.
Additional Materials: Information on medications, catheters, or devices used beyond the routine administration of contrast agents.
Patient Reactions: Document any significant patient reactions or complications, along with any therapeutic interventions.
Patient Instructions: Include any instructions given to the patient or responsible parties.

Heading: Findings

Subheadings:
Observations: Clearly document the radiologic findings using precise anatomic, pathologic, and radiologic terminology.
Clarity: Use short, informative, and factual statements without interpretation. Save interpretation for the impression.
Organization: Group related findings together logically. Consider using lists for multiple related observations to enhance readability.

Heading: Potential Limitations

Subheading:
Study Limitations: Concisely state any limitations of the study that might impact the interpretation of the results.

Heading: Clinical Issues

Subheadings:
Clinical Questions: Address specific clinical questions or issues raised in the referral.
Diagnostic Challenges: Explicitly state if any factors prevent a clear answer to the clinical question.

Heading: Comparison Studies and Reports

Subheading:
Comparative Analysis: Include a comparison with previous studies or reports when available and relevant.

Heading: Impression (Conclusion or Diagnosis)

Subheadings:
Diagnosis: Start with the most likely diagnosis or a differential diagnosis.
Summary: Provide a high-level summary of the key imaging findings that support the diagnosis.
Recommendations: Offer clear, actionable recommendations based on the findings, avoiding unnecessary technical language.
Next Steps: Suggest any follow-up studies or additional diagnostics if necessary.

You must strictly  follow "Rules for Report Generation" for generating the report in json format.

##Rules for Report Generation##
1.	You must only use the information provided in the unstructured report by the radiologist. You must not infer or add details.
2.	You must only include headings and subheadings relevant to the provided information. You must not mention information regarding a heading or subheading if it isn't provided in the unstructured report by the radiologist.
3. If any headings( except 'Demographics','Patient Information' or 'Examination details') or any of their subheadings have empty fields, then you must not mention them in the final output.
4. If subheadings for any of the following headings: Demographics','Patient Information' or 'Examination details' , don't have any content provided, you must leave the space empty in front of the subheadings.
5.	You must ensure that the report is free of unnecessary jargon, especially in the impression, to make it accessible to all healthcare providers.
6.	You must present the information in a logical and chronological order where applicable.

Here is the unstructured report of the radiologist:
    {query}\n'''
    
        output=await gpt_json(create_prompt,1000)
        # print(output)
        report_json=json.loads(output)
        patient_id_exists=True
        result=search_key_in_json(report_json,'patient Id')
        if not result:
            patient_id_exists=False
            patient_info = report_json['Patient Information']
            new_patient_id=await get_latest_patient_id()
            print('new_patient_id',new_patient_id)
            new_patient_info = {'Patient Id': new_patient_id}  # Replace the patient id with the incremented patient id
            new_patient_info.update({k: v for k, v in patient_info.items() if k != 'Patient Id'})
            # Update the original dictionary
            report_json['Patient Information'] = new_patient_info
            json.dumps(report_json, indent=4)
            
            
        print('patient id exists',patient_id_exists)
        #Convert report_json to report, which is some kind of text format of json, but looks human readable
        report = json_to_formatted_string(report_json)
        print("report:\n", report_json)
        return (report,patient_id_exists)
    
    elif intent.lower()=='read':
        read_prompt=f'''You are an AI designed to help doctors to automate Electronic Health Records, you will be given a query by a Doctor, where you are asked to read a medical record for a patient.
         The doctor wants you to first search for a patient based on a specifc attribute(like Name, Age, Surgical Procedure,etc) and the value of that attribute(like Name is Arjun, so Name is the attribute and Arjun is the value) in the Electronic Health Records, and retrive his data. After retrieval the doctor wants you to perfom a task on the retrieved record.
         
         Examples of tasks are: Retrieve Blood Pressure of the patient, Read glucose levels,etc.
         
         You output must be in json format as follows:
         value:<value of the attribute mentioned by the doctor>
         task: <task required to be perfomed on the data, must not contain the Name or value of the attribute.>
         This is the query given by the doctor: 
         {query}\n'''
        #  Name:<search attribute mentioned by the doctor>
        output=await gpt_json(read_prompt,100)
        output=json.loads(output)
        print('output json',output)
        search_attribute = output.get('value')
        task = output.get('task')
        print('search attribute mongodb',search_attribute)
        # Perform text search in MongoDB
        search_query = {"$text": {"$search": search_attribute}}
        cursor = ehr_collection.find(search_query)
        results = []
        patient_ids = set()
        print('cursor',cursor)
        async for document in cursor:
            json_doc = ehr_helper(document)
            results.append(json_doc)
            patient_ids.add(json_doc['patient_id'])
        print('patient ids',patient_ids)
        print('patient ids length',len(patient_ids))
        print(results)
        # Check if all patient IDs are the same
        # if len(patient_ids) > 1:
        #     raise HTTPException(status_code=400, detail="Multiple patient IDs correspond to the search query")
        
        # # Only a single patient id is present
        # # Get all the documents corresponding to the same patient id
        # patient_id=list(patient_ids)[0]
        # cursor=ehr_collection.find({'patient_id': patient_id})
        # async for document in cursor:
        #     json_doc=ehr_helper(document)
        #     results.append(json_doc)
        
        output_of_task=await task_on_EHR(results,datetime.datetime.now(),task)
        #print('output_of_task',output_of_task)
        result_json=search_key_in_json(output_of_task,'output')
        print('output_json_output',result_json)
        outputs=[]
        if isinstance(result_json,list):
            for json_element in result_json:
                # output=json_to_formatted_string(json_element)
                # outputs.append(output)
                if isinstance(json_element, dict):
                    output = json_to_formatted_string(json_element)
                else:
                    output = json_element 
                outputs.append(output)
            output='\n'.join(outputs)
            print('gpt output',output)
            print('out 1')
        else:
            # output=json_to_formatted_string(result_json)
            if isinstance(result_json, dict):
                output = json_to_formatted_string(result_json)
                print('out 2')
            else:
                output = result_json  # Pass through if it's not a dictionary
                print('out 3')
            outputs.append(output)
        #print('output',output)
        return output 
    # elif intent.lower()=='update': 
    #     update_prompt=f'''You are an AI designed to help doctors to automate Electronic Health Records, you will be given a query by a 
    #      Doctor, where you are asked to update the medical record of an existing patient. You need to extract information from the query and structure it into json format with various headings and subheadings that occur in a medical record, like:
    #     Patient details(Name, Patient id, age,etc), Test Reports and Results, Medical History, Prescription, Condition Improvements, Checks and Follow up,etc.
    #      You must keep the Patient id field blank, in case it is not provided.
    #      You must strictly use only the information provided in the query.
    #      You must not mention a Heading or Subheading in the output, if its information isnt given in the query.
    #      You must clearly mention numerical results of the test, and you must structure the report chronologically.
    # This is the query given by the doctor: 
    # {query}\n'''
    #     output=await gpt_json(update_prompt,300)
    #     print('json output',output)
    #     report_json=json.loads(output)
    #     patient_id_exists=True
    #     result=search_key_in_json(report_json,'patient id')
    #     if not result:
    #         patient_id_exists=False
    #         new_report_json = {
    #             "Patient id": result
    #         }
    #         for key, value in report_json.items():
    #             if key.lower() != "patient details":
    #                 new_report_json[key] = value
    #     else:
    #         new_report_json = {
    #             "Patient id": result
    #         }
    #         for key, value in report_json.items():
    #             if key.lower() != "patient details":
    #                 new_report_json[key] = value
                    
    #     print('patient id exists',patient_id_exists)
    #     print('new report',new_report_json)
    #     report = json_to_formatted_string(new_report_json)
    #     print("report:\n", report)
        
    #     return (report,patient_id_exists)# need to break out the actual output
    else:
        return ('No predefined intents match')
        

async def extract_id_and_json_report(report:str):
    summary_prompt = f'''You are an AI designed to help doctors to automate Electronic Health Records, you will be provided with the medical record of a patient as input, and you must convert that medical record to json format.
    You must write the keys in the json according to the headings and subheadins in the record.
    You must extract the Patient id from the medical record,if its provided and give that as output separately.
    If Patient id is not provided in the medical record, you must return an empty string in the output.
    Your output must be in json format as follows:
    Patient id: <patient id>
    Report :<report in json format>
    
    The medical report of the patient:
    {report}'''
    output=await gpt_json(summary_prompt,1000)
    #print('id and json report', output)
    return output

@app.post("/process_request/")
async def process_request(request: Request):
    data = await request.json()
    query = data['text'].strip()
    #print(query)
    #Intent classification here
    intent_prompt=f'''You are an AI designed to help doctors automate Electronic Health Records (EHR). You will be given a query by a doctor, which could involve creating a medical record for a patient or reading existing medical records for a patient.

Your task is to classify the intent of the query into one of the following categories: Create or Read.

You must Use the following criteria to classify the intent:

- **Create:** The query asks to create a new medical record, or modify or add or update information to an existing record.
  Example 1: "Create a new record for John Doe with the diagnosis of hypertension."
  Example 2: "Patient Mathew simons, has reported back with the following test results. add this to the record."

- **Read:** The query asks to retrieve or read existing medical records without making any changes.
  Example: "Retrieve the medical history for Jane Smith."

In your output, you must write the intent clearly as:
Intent: <intent of the query>

This is the query given by the doctor:
{query}\n
    '''
    
    output_intent=await gpt_json(intent_prompt,100)
    output_intent=json.loads(output_intent)
    intent=output_intent["Intent"]
    print("intent",intent)
    output_task=await extract_intent_and_content(query,intent)
    if intent.lower()=='create':
        #send the report to streamlit for the user to edit
        return {"intent": intent,"create_output": output_task[0],"id_exists":output_task[1]}
    elif intent.lower()=='read':
        return {"intent": intent, "read_output":output_task}
    # elif intent.lower()=='update':
    #     return {"intent": intent,"update_output": output_task[0],"id_exists":output_task[1]}
    else:
        return{"intent":'undefined'}

@app.post("/save_request/")
async def save_request(request: Request):
    data = await request.json()
    print(data)
    query = data['text'].strip()
    intent=data.get('intent','')
    print(query)
    print(intent)
    isInvestigation=data.get('isInvestigation',False)
    print('save_request isInvestigation ',isInvestigation)
    #Patient id will be provided-no case where it wont be there
    output_json=await extract_id_and_json_report(query)
    output=json.loads(output_json)
    patient_id=output['Patient id']
    print('patient id gpt',patient_id)
    report=output['Report']
    print('report json gpt',report)
    #Remove Patient Id field from inside the report
    # if 'Patient Id' in report['Patient Information']:
    #     del report['Patient Information']['Patient Id']
    delete_key_from_json(report,'Patient Id')
    entry=1
    entry_investigation=1
    
    #Read the report and find patient id
    patient_id_exists_in_records=False
    if patient_id:
        record_count = await get_patient_record_count(patient_id)
        if record_count!=0:
            patient_id_exists_in_records=True
            entry=record_count+1
        if isInvestigation:
            record_count_inv = await get_patient_record_count(patient_id,isInvestigation=True)
            if record_count_inv !=0:
                entry_investigation=record_count_inv+1
    else:
        return {'message':'Patient id is not provided'}
    
    if patient_id_exists_in_records:
        print('patient id exists in records')
        if isInvestigation:
            updated_record=await add_record(entry_investigation,patient_id,report,isInvestigation=True)
            await print_word_report(patient_id,isInvestigation=True)
            
        updated_record = await add_record(entry, patient_id, report)
        print(updated_record)
        await print_word_report(patient_id)
        return {'message':f' Existing Patient Report has been updated and saved at {patient_id}_report.docx'}
    else:
        print('Patient id doens\'t exist in records, creating a new report if intent is create')
        print(intent)
        if intent.lower()=='update':
            print('inside the condition')
            return {'message':'Patient id does not exist in the database'}
        if isInvestigation:
            updated_record=await add_record(entry_investigation,patient_id,report,isInvestigation=True)
            await print_word_report(patient_id,isInvestigation=True)
            
        created_record = await add_record(entry, patient_id, report)
        print(created_record)
        
        await print_word_report(patient_id)
        
        return {'message':f' A new Patient Report has been created and saved at {patient_id}_report.docx'}

@app.get("/reports")
async def list_reports():
    files = os.listdir(REPORTS_DIR)
    files_with_time = [(file, os.path.getmtime(os.path.join(REPORTS_DIR, file))) for file in files]
    sorted_files = sorted(files_with_time, key=lambda x: x[1], reverse=True)
    recent_files = [file for file, _ in sorted_files[:5]]
    print(recent_files)
    patient_ids=[file.split('_')[0] for file in recent_files]
    patient_names=[]
    for patient_id in patient_ids:
        patient_entry = await patient_id_name_collection.find_one({'patient_id': patient_id})
        if patient_entry:
            # If a matching document is found, return the name
            name= patient_entry.get('name')
        else:
            # If no matching document is found, return None or a default message
            name=''
        patient_names.append(name)
    print(patient_names)
    return {"reports": recent_files,"patient_names":patient_names}

@app.get("/reports_all")
async def list_reports():
    files = os.listdir(REPORTS_DIR)
    patient_ids=[file.split('_')[0] for file in files]
    patient_names=[]
    for patient_id in patient_ids:
        patient_entry = await patient_id_name_collection.find_one({'patient_id': patient_id})
        if patient_entry:
            # If a matching document is found, return the name
            name= patient_entry.get('name')
        else:
            # If no matching document is found, return None or a default message
            name=''
        patient_names.append(name)
    return {"reports": files,"patient_names":patient_names}

@app.get("/reports/{report_name}")
async def get_report(report_name: str):
    report_path = os.path.join(REPORTS_DIR, report_name)
    if os.path.exists(report_path):
        if report_name.endswith(".docx"):
            media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            # Handle PDF or other formats
            media_type = 'application/pdf'
        
        return FileResponse(report_path, media_type=media_type, filename=report_name, headers={
            "Content-Disposition": f"attachment; filename={report_name}"
        })   
    else:
        return {"error": "Report not found"}

@app.post("/filter_reports")
async def filter_reports(request: Request):
    data = await request.json()
    query = data['text'].strip()
    print('query',query)
    query_split=query.split(',')
    print(query_split)
    isInvestigation=data['isInvestigation']
    print('investigation',isInvestigation)
    # search_attribute=''
    # for item in query_split:
    #     search_attribute+=item
    #     search_attribute+=" "
    # print(search_attribute)
    if not query and not isInvestigation:
        files = os.listdir(REPORTS_DIR)
        patient_ids=[file.split('_')[0] for file in files]
        patient_names=[]
        for patient_id in patient_ids:
            patient_entry = await patient_id_name_collection.find_one({'patient_id': patient_id})
            if patient_entry:
                # If a matching document is found, return the name
                name= patient_entry.get('name')
            else:
                # If no matching document is found, return None or a default message
                name=''
            patient_names.append(name)
        return {"reports": files,"patient_names":patient_names}
    
    elif not query and isInvestigation:
        files = os.listdir(REPORTS_DIR)
        filtered_files_inv=[]
        for file in files:
            file_inv=file.split('_')[1]
            if file_inv.split('.')[0]=='investigation':
                filtered_files_inv.append(file)
                
        patient_ids=[file.split('_')[0] for file in filtered_files_inv]
        patient_names=[]
        for patient_id in patient_ids:
            patient_entry = await patient_id_name_collection.find_one({'patient_id': patient_id})
            if patient_entry:
                # If a matching document is found, return the name
                name= patient_entry.get('name')
            else:
                # If no matching document is found, return None or a default message
                name=''
            patient_names.append(name)
        return {"reports": filtered_files_inv,"patient_names":patient_names}
    
    elif query and isInvestigation:
        search_attribute = ' '.join([f"\"{item.strip()}\"" for item in query_split])
        print(search_attribute)
        search_query = {"$text": {"$search": search_attribute}}
        cursor = investigation_collection.find(search_query)
        results = []
        patient_ids = set()
        print('cursor',cursor)
        async for document in cursor:
            json_doc = ehr_helper(document)
            results.append(json_doc)
            patient_ids.add(json_doc['patient_id'])
        print('patient ids',patient_ids)
        print('patient ids length',len(patient_ids))
        #print(results)
        files = os.listdir(REPORTS_DIR)
        filtered_files_inv=[]
        for file in files:
            file_inv=file.split('_')[1]
            if file_inv.split('.')[0]=='investigation':
                filtered_files_inv.append(file)
        #print(files)
        #Logic for getting filtered files- this can change if file name doesnt contain patient id
        filtered_files=[]
        for file in filtered_files_inv:
            file_id=file.split('_')
            if file_id[0] in patient_ids:
                filtered_files.append(file)
        print(filtered_files)
        patient_ids=[file.split('_')[0] for file in filtered_files]
        patient_names=[]
        for patient_id in patient_ids:
            patient_entry = await patient_id_name_collection.find_one({'patient_id': patient_id})
            if patient_entry:
                # If a matching document is found, return the name
                name= patient_entry.get('name')
            else:
                # If no matching document is found, return None or a default message
                name=''
            patient_names.append(name)
        return {"reports": filtered_files,"patient_names":patient_names}
        
    elif query and not isInvestigation:
        search_attribute = ' '.join([f"\"{item.strip()}\"" for item in query_split])
        print(search_attribute)
        search_query = {"$text": {"$search": search_attribute}}
        cursor = ehr_collection.find(search_query)
        results = []
        patient_ids = set()
        print('cursor',cursor)
        async for document in cursor:
            json_doc = ehr_helper(document)
            results.append(json_doc)
            patient_ids.add(json_doc['patient_id'])
        print('patient ids',patient_ids)
        print('patient ids length',len(patient_ids))
        print(results)
        files = os.listdir(REPORTS_DIR)
        print(files)
        #Logic for getting filtered files- this can change if file name doesnt contain patient id
        filtered_files=[]
        for file in files:
            file_id=file.split('_')
            if file_id[0] in patient_ids:
                filtered_files.append(file)
        print(filtered_files)
        
        patient_ids=[file.split('_')[0] for file in filtered_files]
        patient_names=[]
        for patient_id in patient_ids:
            patient_entry = await patient_id_name_collection.find_one({'patient_id': patient_id})
            if patient_entry:
                # If a matching document is found, return the name
                name= patient_entry.get('name')
            else:
                # If no matching document is found, return None or a default message
                name=''
            patient_names.append(name)
        return {"reports": filtered_files,"patient_names":patient_names}
        

@app.post('/process_file')
async def process_file(file: UploadFile = File(...)):
    try:
        # Read the content of the uploaded file
        content = await file.read()
        if not content:
            print("No content read from the file")
            return {"error": "No content read from the file"}

        print("File content read successfully")

        text = ""
        
        # Convert PDF pages to images
        images = convert_from_bytes(content)
        if not images:
            print("No images extracted from the PDF")
            return {"error": "No images extracted from the PDF"}
        
        print(f"{len(images)} pages extracted from the PDF")

        # Perform OCR on each image
        for page_num, image in enumerate(images):
            print(f'Processing page {page_num + 1}')
            page_text = pytesseract.image_to_string(image)
            if page_text:
                text += page_text + "\n"
                print(f"Extracted text from page {page_num + 1}: {page_text}")
            else:
                print(f"No text extracted from page {page_num + 1}")
        
        if not text:
            print("No text extracted from the entire PDF")
            return {"error": "No text extracted from the entire PDF"}
        
        print(f"Full OCR Text: {text}")
        #Some kind of report preprocessing is required, this can change with time.
        report_mod_prompt=f'''You are an AI designed to help doctors automate Electronic Health Records (EHR). You will be provided with the extracted text from a patient report. This text will contain information regarding the patient, and his medical condition and other relevant details. The text may also contain details like report headers which may include hospital name, location,etc.
        Your task is to identify the relevant medical report of the patient, and extract it.You must also search for a Patient id field in the report. In case the patient id field is not present. You must write a blank field of 'Patient id: ' , in your output report.
        Your output must be in a json format as follows:
        medical_report: <The extracted medical report>
        
        The input report:
        {text}
        '''
        output_report=await gpt_json(report_mod_prompt,300)
        output_report=json.loads(output_report)
        report = json_to_formatted_string(output_report['medical_report'])
        return {"text": report}
    except Exception as e:
        print(f"Exception: {str(e)}")
        return {"error": str(e)}

    
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)